from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CANONICAL_DIR = ROOT / "canonical_chapters"
FORMAL_DIR = ROOT / "formal_docx"
FORMAT_DIR = ROOT / "format"
FORMAL_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = FORMAL_DIR / "中国人民大学本科毕业论文_正式排版版.docx"
ASSET_DIR = FORMAL_DIR / "assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
REFERENCES_PATH = CANONICAL_DIR / "references.md"
FORMAT_LOGO_PATH = FORMAT_DIR / "logo.png"
PRIMARY_LOGO_PATH = FORMAT_LOGO_PATH if FORMAT_LOGO_PATH.exists() else ASSET_DIR / "image1.png"
HEADER_LOGO_PATH = FORMAT_LOGO_PATH if FORMAT_LOGO_PATH.exists() else ASSET_DIR / "image2.jpeg"

TITLE = "AI 技术事件与职业不安全感表达的关联"
SUBTITLE = "——基于社交媒体数据的描述性分析"
AUTHOR = "石盛楠"
MAJOR = "人力资源管理"
COLLEGE = "中法学院"
GRADE = "2022级"
STUDENT_ID = "2022202523"
ADVISOR = ""
PAPER_CODE = "RUC-BK-待填写"
PAPER_SCORE = ""
DATE_TEXT = "2026年4月"

CN_ABSTRACT = (
    "人工智能快速发展引发了公众对岗位替代与职业安全的持续讨论。本文基于知乎、哔哩哔哩和小红书三个平台 "
    "2024 年 10 月至 2026 年 3 月的 114,915 条评论，采用词典识别方法提取 1,242 条职业不安全感相关表达，"
    "并结合 33 个重大 AI 技术事件，使用描述统计、负二项回归、事件研究和趋势分析考察其表达模式。结果显示："
    "知乎表达率最高（2.80%），哔哩哔哩次之（0.86%），小红书因样本过少仅作描述；相关表达总体随时间显著上升"
    "（IRR=1.030, p<0.001），事件窗口主效应不显著（IRR=0.653, p=0.083），事件后的相关表达约在 3 天内回落至周边均值。"
    "研究表明，社交媒体可用于观察职业不安全感相关表达的跨平台差异与时间变化，但不能直接等同于职业不安全感构念水平，"
    "也不能据此进行因果推断。"
)

CN_ABSTRACT_2 = ""

CN_KEYWORDS = "人工智能；职业不安全感；社交媒体分析；平台差异；时间趋势"

EN_ABSTRACT = (
    "The rapid development of artificial intelligence has intensified public discussion about job displacement and "
    "occupational security. Existing studies on AI-related threat perception still rely heavily on cross-sectional "
    "surveys and therefore provide limited evidence about how public responses evolve in natural settings. "
    "This thesis examines the association between AI technology events and job insecurity expression by analyzing "
    "114,915 comments collected from Zhihu, Bilibili, and Xiaohongshu between October 2024 and March 2026. "
    "A dictionary-based identification scheme was used to detect 1,242 job-insecurity-related expressions, and an "
    "event database covering 33 major AI technology events was constructed. Descriptive statistics, negative "
    "binomial regression, event-study analysis, and time-trend analysis were then applied under a natural-day "
    "aggregation framework."
)

EN_ABSTRACT_2 = (
    "The results show clear platform differences in observed expression rates, with Zhihu showing the highest rate "
    "(2.80%) and Bilibili the second highest (0.86%). Because Xiaohongshu covers only three observed dates, it is "
    "reported descriptively rather than inferentially. Job insecurity expression also displays a significant upward "
    "time trend during the study period (IRR = 1.030, p < 0.001). By contrast, the main event-window effect does not "
    "reach conventional significance (IRR = 0.653, p = 0.083), indicating that the short-term effect of AI events is "
    "not robustly positive. Event-study results further suggest that observed expression returns to the surrounding "
    "mean level within about three days after an event. Overall, the thesis identifies patterns of public expression, "
    "platform variation, and temporal change in open social-media discussions, but it does not directly measure the "
    "latent construct of job insecurity or support causal inference."
)

EN_KEYWORDS = "artificial intelligence; job insecurity; social media analysis; platform differences; time trend"


def ensure_styles(document: Document) -> None:
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        _ = document.styles[style_name]

    if "TOC Heading" not in document.styles:
        document.styles.add_style("TOC Heading", WD_STYLE_TYPE.PARAGRAPH)


def set_run_fonts(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman") -> None:
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:cs"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def format_run(run, *, east_asia: str = "宋体", ascii_font: str = "Times New Roman",
               size: float = 12, bold: bool = False, italic: bool = False,
               underline: bool = False) -> None:
    set_run_fonts(run, east_asia=east_asia, ascii_font=ascii_font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def configure_section(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    sect_pr = section._sectPr
    pg_mar = sect_pr.pgMar
    if pg_mar is not None:
        pg_mar.set(qn("w:gutter"), str(int(Cm(0.5).emu / 635)))


def set_page_numbering(section, *, start: int = 1, fmt: str = "decimal") -> None:
    sect_pr = section._sectPr
    for existing in sect_pr.xpath("./w:pgNumType"):
        sect_pr.remove(existing)
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), fmt)
    sect_pr.append(pg_num)


def ensure_logo_assets() -> None:
    if PRIMARY_LOGO_PATH.exists() and HEADER_LOGO_PATH.exists():
        return
    docx_files = sorted(FORMAT_DIR.glob("*.docx"))
    if not docx_files:
        return
    with zipfile.ZipFile(docx_files[0]) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/") or name.endswith("/"):
                continue
            target = ASSET_DIR / Path(name).name
            if not target.exists():
                target.write_bytes(archive.read(name))


def add_update_fields_on_open(document: Document) -> None:
    settings_part = document.settings.element
    update_fields = settings_part.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_part.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_page_field(paragraph, instruction: str, displayed: str = "1") -> None:
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    run_instr._r.append(instr_text)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    run_text = paragraph.add_run(displayed)
    format_run(run_text, east_asia="宋体", ascii_font="Times New Roman", size=10.5)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_picture_to_paragraph(paragraph, image_path: Path, width_cm: float, height_cm: float | None = None) -> bool:
    if not image_path.exists():
        return False
    run = paragraph.add_run()
    if height_cm is None:
        run.add_picture(str(image_path), width=Cm(width_cm))
    else:
        run.add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
    return True


def set_header(section, text: str) -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    if not add_picture_to_paragraph(paragraph, HEADER_LOGO_PATH, 4.13, 0.98):
        run = paragraph.add_run(text)
        format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5)


def clear_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()


def set_footer(section, *, roman: bool) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    if roman:
        add_page_field(paragraph, "PAGE", "I")
    else:
        prefix = paragraph.add_run("第")
        format_run(prefix, east_asia="宋体", ascii_font="Times New Roman", size=10.5)
        add_page_field(paragraph, "PAGE", "1")
        suffix = paragraph.add_run("页")
        format_run(suffix, east_asia="宋体", ascii_font="Times New Roman", size=10.5)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    code_value = PAPER_CODE if PAPER_CODE.strip() and "待填写" not in PAPER_CODE else "______________"
    run = p.add_run(f"论文编码：{code_value}")
    format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=12, bold=True)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(14)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国人民大学本科毕业论文（设计）")
    format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=28, bold=True)

    for text in [TITLE, SUBTITLE]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=28 if text == TITLE else 20, bold=True)

    for _ in range(4):
        document.add_paragraph()

    cover_fields = [
        ("作者", AUTHOR),
        ("学院", COLLEGE),
        ("专业", MAJOR),
        ("年级", GRADE),
        ("学号", STUDENT_ID),
        ("指导教师", ADVISOR),
        ("论文成绩", PAPER_SCORE),
        ("日期", DATE_TEXT),
    ]
    for label, value in cover_fields:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Pt(100)
        p.paragraph_format.space_after = Pt(6)
        run_label = p.add_run(f"{label}：")
        format_run(run_label, east_asia="黑体", ascii_font="Times New Roman", size=20, bold=True)
        display_value = value if value.strip() else " " * 18
        run_value = p.add_run(display_value)
        format_run(run_value, east_asia="宋体", ascii_font="Times New Roman", size=20, underline=True)


def add_declaration_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国人民大学学位论文原创性声明和使用授权说明")
    format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)

    for text in [
        "本人郑重声明：所呈交的论文是我个人在导师指导下进行的研究工作及取得的研究成果。尽我所知，除了文中特别加以标注和致谢的地方外，论文中不包含其他人已经发表或撰写过的研究成果，也不包含为获得中国人民大学或其他教育机构的学位或证书所使用过的材料。与我一同工作的同志对本研究所做的任何贡献均已在文中作了明确的说明并表示了谢意。",
        "本人完全了解中国人民大学有关保留、使用学位论文的规定，即：学校有权保留送交论文的复印件，允许论文被查阅和借阅；学校可以公布论文的全部或部分内容，可以采用影印、缩印或其他复制手段保存论文。",
    ]:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(text)
        format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)

    for label in [
        "作者签名：____________________",
        "日期：____________________",
        "指导教师签名：____________________",
        "日期：____________________",
    ]:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(label)
        format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def add_abstract_page(document: Document, *, english: bool) -> None:
    title = "ABSTRACT" if english else "摘  要"
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    format_run(
        run,
        east_asia="黑体" if not english else "Times New Roman",
        ascii_font="Times New Roman",
        size=16,
        bold=True,
    )
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    body_paragraphs = [EN_ABSTRACT, EN_ABSTRACT_2] if english else [CN_ABSTRACT, CN_ABSTRACT_2]
    for para_text in body_paragraphs:
        if not para_text.strip():
            continue
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 2 if english else 1.5
        run = p.add_run(para_text)
        format_run(
            run,
            east_asia="宋体" if not english else "Times New Roman",
            ascii_font="Times New Roman",
            size=12,
        )

    p = document.add_paragraph()
    label = "Key Words：" if english else "关键词："
    value = EN_KEYWORDS if english else CN_KEYWORDS
    p.paragraph_format.space_before = Pt(12)
    label_run = p.add_run(label)
    format_run(
        label_run,
        east_asia="黑体" if not english else "Times New Roman",
        ascii_font="Times New Roman",
        size=12,
        bold=True,
    )
    value_run = p.add_run(value)
    format_run(value_run, east_asia="宋体" if not english else "Times New Roman", ascii_font="Times New Roman", size=12)


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目  录")
    format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_page_field(p, 'TOC \\o "1-3" \\h \\z \\u', "目录将在打开文档后自动更新")


def extract_table_captions() -> list[str]:
    captions: list[str] = []
    for path in sorted(CANONICAL_DIR.glob("*.md")):
        if path.name in {"references.md", "final_consistency_matrix.md", "citations_extracted.txt"}:
            continue
        for raw_line in path.read_text(encoding="utf-8").splitlines():
            text = clean_inline_markdown(raw_line.strip())
            if re.match(r"^表\s*\d+\s+", text):
                captions.append(text)
    return captions


def add_table_directory(document: Document) -> None:
    captions = extract_table_captions()
    if not captions:
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("表目录")
    format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)

    for caption in captions:
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(caption)
        format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace("\\", "")
    return text.strip()


def format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.25
    for run in paragraph.runs:
        format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def format_heading(paragraph, level: int) -> None:
    sizes = {1: 16, 2: 14, 3: 12, 4: 10.5}
    paragraph.style = f"Heading {min(level, 4)}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    for run in paragraph.runs:
        format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=sizes[min(level, 4)], bold=True)


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$", line))


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [clean_inline_markdown(cell.strip()) for cell in stripped.split("|")]


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.0
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=10.5, bold=(r_idx == 0))


def add_table_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=True)


def add_list_paragraph(document: Document, text: str, numbered: bool = False) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-12)
    p.paragraph_format.line_spacing = 1.25
    prefix = "1. " if numbered else "• "
    run = p.add_run(prefix + clean_inline_markdown(text))
    format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def structural_start(line: str) -> bool:
    stripped = line.strip()
    return (
        stripped.startswith("#")
        or stripped.startswith("- ")
        or bool(re.match(r"^\d+\.\s", stripped))
        or stripped.startswith("|")
        or stripped.startswith("$$")
        or stripped == "---"
    )


def add_markdown_file(document: Document, path: Path) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = clean_inline_markdown(stripped[level:].strip())
            p = document.add_paragraph()
            p.add_run(text)
            format_heading(p, level)
            i += 1
            continue
        cleaned = clean_inline_markdown(stripped)
        if re.match(r"^表\s*\d+\s+", cleaned):
            add_table_caption(document, cleaned)
            i += 1
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(split_table_row(lines[i].strip()))
                i += 1
            add_markdown_table(document, table_lines)
            continue
        if stripped.startswith("- "):
            add_list_paragraph(document, stripped[2:])
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            add_list_paragraph(document, re.sub(r"^\d+\.\s*", "", stripped), numbered=True)
            i += 1
            continue
        if stripped.startswith("$$") and stripped.endswith("$$"):
            text = stripped.strip("$")
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)
            i += 1
            continue
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line or structural_start(next_line):
                break
            para_lines.append(next_line)
            i += 1
        paragraph = document.add_paragraph(clean_inline_markdown(" ".join(para_lines)))
        format_body_paragraph(paragraph)


def iter_reference_entries() -> list[str]:
    if not REFERENCES_PATH.exists():
        return []
    entries: list[str] = []
    skip_section = False
    for raw_line in REFERENCES_PATH.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line == "---" or line.startswith("# ") or line.startswith(">"):
            continue
        if line.startswith("## "):
            heading = clean_inline_markdown(line[3:].strip())
            skip_section = "已删除" in heading or "虚构" in heading
            continue
        if skip_section:
            continue
        entries.append(clean_inline_markdown(line))
    return entries


def normalize_reference_entry(text: str) -> str:
    text = text.strip().rstrip(".")
    match = re.match(r"^(?P<authors>.+?)\s+\((?P<year>\d{4})\)\.\s+(?P<rest>.+)$", text)
    if not match:
        return text

    authors = match.group("authors").strip().rstrip(".")
    year = match.group("year").strip()
    rest = match.group("rest").strip()

    url_match = re.search(r"(https?://\S+)$", rest)
    url = ""
    if url_match:
        url = url_match.group(1)
        rest = rest[:url_match.start()].rstrip().rstrip(".")

    suffix = f" {url}" if url else ""

    if ". In " in rest:
        title, container = rest.split(". In ", 1)
        pages_match = re.search(r"\(pp\.\s*([^)]+)\)", container)
        pages = pages_match.group(1) if pages_match else ""
        container_no_pages = re.sub(r"\s*\(pp\.\s*([^)]+)\)", "", container).strip().rstrip(".")
        return f"{authors}. {title}[A]. In: {container_no_pages}[C]. {year}: {pages}。{suffix}".rstrip()

    journal_match = re.match(
        r"^(?P<title>.+?)\.\s+(?P<journal>.+?),\s+(?P<vol>\d+)(?:\((?P<issue>[^)]+)\))?(?:,\s+(?P<pages>.+))?$",
        rest,
    )
    if journal_match:
        title = journal_match.group("title").strip()
        journal = journal_match.group("journal").strip()
        vol = journal_match.group("vol").strip()
        issue = journal_match.group("issue")
        pages = (journal_match.group("pages") or "").strip()
        issue_text = f"({issue})" if issue else ""
        page_text = f": {pages}" if pages else ""
        return f"{authors}. {title}[J]. {journal}, {year}, {vol}{issue_text}{page_text}。{suffix}".rstrip()

    if "arXiv preprint" in rest:
        title, source = rest.split(". ", 1)
        return f"{authors}. {title}[EB/OL]. {source}, {year}。{suffix}".rstrip()

    report_keywords = ["Working Papers", "Institute", "Forum", "OECD", "McKinsey", "Report", "Publishing"]
    if any(keyword in rest for keyword in report_keywords):
        title, source = rest.split(". ", 1) if ". " in rest else (rest, "")
        source_text = f" {source}" if source else ""
        return f"{authors}. {title}[R].{source_text} {year}。{suffix}".rstrip()

    if any(keyword in rest for keyword in ["Press", "Publishing", "Sons", "Institute", "University of Chicago", "OECD"]):
        title, publisher = rest.split(". ", 1) if ". " in rest else (rest, "")
        publisher_text = f" {publisher}" if publisher else ""
        return f"{authors}. {title}[M].{publisher_text} {year}。{suffix}".rstrip()

    return f"{authors}. {rest}. {year}。{suffix}".rstrip()


def add_references_section(document: Document) -> None:
    p = document.add_paragraph()
    p.add_run("参考文献")
    format_heading(p, 1)

    entries = iter_reference_entries()
    if not entries:
        p = document.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run("当前未找到 references.md，参考文献页保留为空，请在定稿前补齐。")
        format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)
        return

    for idx, text in enumerate(entries, start=1):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(f"[{idx}] {normalize_reference_entry(text)}")
        format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def add_signature_section(document: Document) -> None:
    p = document.add_paragraph()
    run = p.add_run("作者签名：____________________")
    format_run(run, east_asia="黑体", ascii_font="Times New Roman", size=14, bold=True)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)


def add_acknowledgments(document: Document) -> None:
    p = document.add_paragraph()
    p.add_run("致谢")
    format_heading(p, 1)

    p = document.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(
        "在论文完成过程中，笔者得到学院教师、课程老师与同学们的帮助和支持。"
        "谨向所有在选题讨论、资料整理、写作修改和研究推进过程中给予关心与建议的师长和同学表示诚挚感谢。"
        "同时，也感谢相关平台公开信息为本研究提供的数据基础。文中如有疏漏与不足，均由笔者本人负责。"
    )
    format_run(run, east_asia="宋体", ascii_font="Times New Roman", size=12)


def main() -> None:
    ensure_logo_assets()
    document = Document()
    ensure_styles(document)
    configure_section(document.sections[0])
    add_update_fields_on_open(document)
    set_header(document.sections[0], "中国人民大学本科毕业论文（设计）")
    clear_footer(document.sections[0])

    add_cover(document)

    front_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front_section)
    set_page_numbering(front_section, start=1, fmt="upperRoman")
    set_header(front_section, "中国人民大学本科毕业论文（设计）")
    set_footer(front_section, roman=True)

    add_declaration_page(document)
    document.add_page_break()
    add_abstract_page(document, english=False)
    document.add_page_break()
    add_abstract_page(document, english=True)
    document.add_page_break()
    add_toc(document)
    document.add_page_break()
    add_table_directory(document)

    body_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section)
    set_page_numbering(body_section, start=1, fmt="decimal")
    set_header(body_section, "中国人民大学本科毕业论文（设计）")
    set_footer(body_section, roman=False)

    for chapter_name in [
        "01_introduction.md",
        "02_literature_review.md",
        "03_methodology.md",
        "04_results.md",
        "05_discussion.md",
        "06_conclusion.md",
    ]:
        add_markdown_file(document, CANONICAL_DIR / chapter_name)

    add_signature_section(document)
    add_references_section(document)
    add_acknowledgments(document)

    document.save(str(OUTPUT_PATH))
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
